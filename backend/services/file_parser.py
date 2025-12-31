"""
智能脚本解析服务 - 使用 AI 解析任意格式的脚本文件
"""
import json
import uuid
import logging
import httpx
from typing import List, Dict, Any
from dataclasses import dataclass, field, asdict
from enum import Enum
from io import BytesIO
import openpyxl
from docx import Document
from config import Config

logger = logging.getLogger(__name__)


class PageStatus(str, Enum):
    """页面状态"""
    PENDING = "pending"
    GENERATING_DESC = "generating_desc"
    GENERATING_IMAGE = "generating_image"
    COMPLETED = "completed"
    ERROR = "error"


@dataclass
class ScriptPage:
    """脚本页面数据"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    index: int = 0
    shot_number: str = ""  # 镜号
    segment: str = ""  # 环节
    narration: str = ""  # 讲稿内容
    visual_hint: str = ""  # 画面描述（可选）
    description: str = ""  # AI 生成的页面描述
    image_path: str = ""  # 生成的图片路径
    status: PageStatus = PageStatus.PENDING
    error_message: str = ""

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        data = asdict(self)
        data['status'] = self.status.value
        return data


def excel_to_text(file_path: str) -> str:
    """将 Excel 文件转换为文本格式，保留表格结构"""
    wb = openpyxl.load_workbook(file_path, read_only=True)
    ws = wb.active

    lines = []
    for row in ws.iter_rows(values_only=True):
        # 过滤空行
        if any(cell is not None and str(cell).strip() for cell in row):
            row_text = ' | '.join(str(cell) if cell else '' for cell in row)
            lines.append(row_text)

    wb.close()
    return '\n'.join(lines)


def docx_to_text(file_path: str) -> str:
    """将 Word 文件转换为文本格式，包括段落和表格"""
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"无法解析 Word 文件 {file_path}: {e}")
        raise ValueError(f"Word 文件解析失败: {e}")

    lines = []
    # 提取段落
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return '\n'.join(lines)


def parse_with_ai(content: str) -> List[Dict[str, Any]]:
    """使用 AI 智能解析脚本内容（Gemini 原生 API）"""

    prompt = f"""你是一个脚本解析助手。请分析以下内容，提取出每一页/每一镜的信息。

内容可能是各种格式：Excel表格、纯文本、分段文字等。请智能识别并提取。

对于每一页，请提取：
- shot_number: 镜号/页码（如果有的话，字符串格式）
- segment: 环节/章节名称（如果有的话）
- narration: 讲稿/旁白/主要文字内容（必须有，这是最重要的内容）
- visual_hint: 画面描述/视觉提示（如果有的话）

请以 JSON 数组格式返回，每个元素是一页的信息。只返回 JSON，不要其他说明。

规则：
1. 如果内容没有明确的页/镜划分，请按照自然段落或逻辑单元拆分
2. 跳过明显是元数据的行（如标题行、作者信息等）
3. 跳过空行或只有"/"的行
4. narration 是核心内容，必须提取

===== 脚本内容 =====
{content}
===== 内容结束 =====

返回格式示例：
[
  {{"shot_number": "1", "segment": "开场", "narration": "大家好...", "visual_hint": "标题画面"}},
  {{"shot_number": "2", "segment": "正文", "narration": "今天我们来学习...", "visual_hint": ""}}
]"""

    # 使用 Gemini 原生 API
    url = f"{Config.TEXT_API_BASE}/v1beta/models/{Config.TEXT_MODEL}:generateContent"
    headers = {
        "x-goog-api-key": Config.TEXT_API_KEY,
        "Content-Type": "application/json"
    }
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.3,
            "maxOutputTokens": 8192
        }
    }

    with httpx.Client(timeout=120.0) as client:
        response = client.post(url, json=payload, headers=headers)
        response.raise_for_status()
        result = response.json()

    # 提取文本响应
    result_text = ""
    if "candidates" in result and len(result["candidates"]) > 0:
        candidate = result["candidates"][0]
        if "content" in candidate and "parts" in candidate["content"]:
            for part in candidate["content"]["parts"]:
                if "text" in part:
                    result_text += part["text"]

    result_text = result_text.strip()

    if not result_text:
        logger.error("[file_parser] API 返回内容为空")
        raise ValueError("AI 返回内容为空")

    # 清理可能的 markdown 代码块标记
    if result_text.startswith('```'):
        lines = result_text.split('\n')
        # 移除首尾的 ``` 行
        if lines[0].startswith('```'):
            lines = lines[1:]
        if lines and lines[-1].strip() == '```':
            lines = lines[:-1]
        result_text = '\n'.join(lines)

    # 过滤思考过程，提取 JSON 数组
    import re
    json_match = re.search(r'\[[\s\S]*\]', result_text)
    if json_match:
        result_text = json_match.group()

    try:
        pages_data = json.loads(result_text)
    except json.JSONDecodeError as e:
        logger.error(f"AI 返回内容无法解析: {result_text[:500]}")
        raise ValueError(f"AI 返回的内容无法解析为 JSON: {e}")

    return pages_data


class FileParser:
    """智能文件解析器"""

    @staticmethod
    def parse_file(file_path: str) -> List[ScriptPage]:
        """
        智能解析脚本文件（任意格式）

        Args:
            file_path: 文件路径

        Returns:
            ScriptPage 列表
        """
        from pathlib import Path
        path = Path(file_path)
        suffix = path.suffix.lower()

        # 转换为文本
        if suffix in ['.xlsx', '.xls']:
            text_content = excel_to_text(file_path)
        elif suffix == '.docx':
            text_content = docx_to_text(file_path)
        elif suffix in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        elif suffix == '.json':
            # JSON 可以直接解析，但也可以让 AI 处理更灵活
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
        else:
            # 尝试当作文本处理
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            except (UnicodeDecodeError, IOError, OSError) as e:
                logger.error(f"无法读取文件 {file_path}: {e}")
                raise ValueError(f"不支持的文件格式: {suffix}")

        logger.info(f"读取文件内容，长度: {len(text_content)} 字符")

        # 使用 AI 解析
        pages_data = parse_with_ai(text_content)
        logger.info(f"AI 解析出 {len(pages_data)} 页")

        # 转换为 ScriptPage 对象
        pages = []
        for i, page_data in enumerate(pages_data):
            narration = page_data.get('narration', '')
            if not narration or narration.strip() in ['', '/']:
                continue

            page = ScriptPage(
                index=i,
                shot_number=str(page_data.get('shot_number', i + 1)),
                segment=page_data.get('segment', ''),
                narration=narration,
                visual_hint=page_data.get('visual_hint', '')
            )
            pages.append(page)

        logger.info(f"最终解析出 {len(pages)} 个有效页面")
        return pages
